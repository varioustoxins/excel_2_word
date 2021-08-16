from collections import OrderedDict
import sys
import xlrd
from docx import Document
from os.path import splitext


POSTER_ABSTRACT = 'Poster abstract'
POSTER_TITLE = 'Poster title'
INSTITUTION = 'Institution'
FAMILY_NAME = 'Family name (surname)'
FIRST_NAMES = 'Given name (first name)'




if __name__ == '__main__':

    if len(sys.argv) < 2:
        print('I need a file to work on!')
        print('Exiting...')
        sys.exit(-1)
    elif len(sys.argv) > 2:
        print('I can only work with 1 file!')
        print('Exiting...')
        sys.exit(-1)
    
    file_name = sys.argv[1]
    
    if not file_name.endswith('.xls'):
       print(f'I need an xls file i got {file_name} is that an xlxs file?')
       print('Exiting...')
       sys.exit(-1)
     
    loc =(file_name)
    root=splitext(loc)[0]
    wb = xlrd.open_workbook(loc)
    sheet = wb.sheet_by_index(0)

    cols = OrderedDict()
    for i in range(sheet.ncols):
        cols[sheet.cell_value(0, i)] = i


    required_cols= [FIRST_NAMES, FAMILY_NAME, INSTITUTION, POSTER_TITLE, POSTER_ABSTRACT]

    document = Document()

    count = 1
    for row in range(1,sheet.nrows):
        title = sheet.cell_value(row, cols[POSTER_TITLE])
        if len(title.strip()) != 0:
            if title  not in ['NA','n/a']:
                first_names =  sheet.cell_value(row, cols[FIRST_NAMES])
                family_name = sheet.cell_value(row, cols[FAMILY_NAME])
                institution = sheet.cell_value(row, cols[INSTITUTION])
                abstract =  sheet.cell_value(row, cols[POSTER_ABSTRACT])

                document.add_heading('%i. %s' % (count, title), level=1)
                document.add_heading('%s %s' % (first_names, family_name), level=2)
                document.add_heading('%s' % institution, level=3)
                document.add_paragraph('')
                document.add_paragraph(abstract)
                document.add_paragraph('')
                document.add_page_break()
                count = count+1


    document.save('%s.docx' % root)
